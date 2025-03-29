import os
import pdfplumber
from docx import Document
from openai import OpenAI, APIConnectionError
import yaml
import tempfile
import subprocess
import io
import time

current_path = os.path.dirname(os.path.abspath(__file__))


def transcribe_audio(audio):
    # Transcribe audio file using OpenAI Whisper
    client = OpenAI()
    attempt = 0
    while attempt < 3:
        try:
            transcription = client.audio.transcriptions.create(
                model="whisper-1", file=audio
            )
            return transcription.text
        except APIConnectionError as e:
            print(f"Attempt {attempt + 1} failed: {e}")
            attempt += 1
            if attempt < 3:
                print("Retrying in 5 seconds...")
                time.sleep(5)

    print("Maximum number of attempts reached. Transcription failed.")
    return None


def extract_text_from_docx(docx) -> str:
    # Load Word document from file name

    doc = Document(docx)
    paragraphs = [p.text for p in doc.paragraphs]
    # Extract paragraphs
    paragraphs = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Extract tables
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            tables_text.append(row_text)

    tables_text = "\n".join(tables_text)

    # Combine content
    result_text = (
        f"{paragraphs}\n\nTables:\n{tables_text}" if tables_text else paragraphs
    )

    return result_text


def extract_text_from_pdf(file) -> str:
    # Load PDF file from file path
    reader = pdfplumber.open(file)
    data = "\n".join(
        [page.extract_text() for page in reader.pages if page.extract_text()]
    )
    return data


def extract_prompts(yaml_dir, model_name):
    current_path = os.path.dirname(os.path.abspath(__file__))
    with open(os.path.join(current_path, yaml_dir), "r") as file:
        prompts = yaml.safe_load(file)
        return prompts[f"report_generation_{model_name}"]


def fill_prompt_template(prompt_template, replacements):
    filled_prompt = prompt_template
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"  # Convert key to {{KEY}} format
        filled_prompt = filled_prompt.replace(placeholder, value)
    return filled_prompt


def compress_audio_ffmpeg(uploaded_file):
    """
    Compress audio using FFmpeg with settings optimized for speech recognition.

    Parameters:
    - uploaded_file: Streamlit UploadedFile object

    Returns:
    - File-like object containing the compressed audio data
    """
    try:
        # Create a temporary file for the input
        input_ext = os.path.splitext(uploaded_file.name)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=input_ext) as input_tmp:
            input_tmp.write(uploaded_file.getvalue())
            input_path = input_tmp.name

        # Create a path for the output file
        output_path = tempfile.mktemp(suffix=".ogg")

        # FFmpeg command to compress audio for voice
        cmd = [
            "ffmpeg",
            "-i",
            input_path,  # Input file
            "-vn",  # No video
            "-map_metadata",
            "-1",  # Remove metadata
            "-ac",
            "1",  # Convert to mono
            "-c:a",
            "libopus",  # Use Opus codec
            "-b:a",
            "12k",  # 12kbps bitrate
            "-application",
            "voip",  # Optimize for voice
            output_path,  # Output file
        ]

        # Run FFmpeg command
        subprocess.run(cmd, check=True, capture_output=True)

        # Clean up the input temporary file
        os.unlink(input_path)

        # Log file sizes
        original_size = len(uploaded_file.getvalue()) / (1024 * 1024)
        compressed_size = os.path.getsize(output_path) / (1024 * 1024)
        print(
            f"Original size: {original_size:.2f}MB, Compressed size: {compressed_size:.2f}MB"
        )

        # Read the compressed file into a BytesIO object
        with open(output_path, "rb") as f:
            compressed_file = io.BytesIO(f.read())

        # Set the name of the file to match the original but with .ogg extension
        original_name = os.path.splitext(uploaded_file.name)[0]
        compressed_file.name = f"{original_name}.ogg"

        # Clean up the output temporary file
        os.unlink(output_path)

        return compressed_file
    except Exception as e:
        print(f"Compression failed: {e}")
        return None
